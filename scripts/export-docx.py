#!/usr/bin/env python3
"""
Schreibwerkstatt — DOCX Export
Kombiniert alle Kapitel zu einem formatierten Word-Dokument.
"""

import os
import re
import sys
from pathlib import Path

def find_project_root():
    """Findet das Projektstammverzeichnis."""
    current = Path.cwd()
    while current != current.parent:
        if (current / "CLAUDE.md").exists() and (current / "bibel").exists():
            return current
        current = current.parent
    return Path.cwd()

def get_chapters(kapitel_dir):
    """Liest alle Kapitel in der richtigen Reihenfolge."""
    chapters = []
    for f in sorted(kapitel_dir.glob("*.md")):
        if f.name.startswith("_"):
            continue
        content = f.read_text(encoding="utf-8")
        lines = content.strip().split("\n")
        title = lines[0].lstrip("# ").strip() if lines else f.stem
        body = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""
        chapters.append({"title": title, "body": body, "file": f.name})
    return chapters

def get_metadata(root):
    """Liest Metadaten aus der Synopsis."""
    synopsis = root / "geschichte" / "synopsis.md"
    metadata = {"title": "Unbenannter Roman", "author": "Unbekannt"}
    if synopsis.exists():
        content = synopsis.read_text(encoding="utf-8")
        title_match = re.search(r"## Arbeitstitel\s*\n+(.+)", content)
        if title_match:
            title = title_match.group(1).strip().strip("<!--").strip("-->").strip()
            if title:
                metadata["title"] = title
    return metadata

def export_docx(root, output_path):
    """Erstellt eine DOCX-Datei."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ python-docx nicht installiert!")
        print("   Installiere mit: pip3 install python-docx")
        sys.exit(1)

    kapitel_dir = root / "_bmad-output" / "kapitel"
    if not kapitel_dir.exists() or not list(kapitel_dir.glob("*.md")):
        print("❌ Keine Kapitel gefunden in _bmad-output/kapitel/")
        sys.exit(1)

    chapters = get_chapters(kapitel_dir)
    metadata = get_metadata(root)

    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Titelseite
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(metadata["title"])
    title_run.font.size = Pt(28)
    title_run.bold = True

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(metadata["author"])
    author_run.font.size = Pt(16)

    doc.add_page_break()

    # Kapitel
    for ch in chapters:
        heading = doc.add_heading(ch["title"], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Absätze verarbeiten
        paragraphs = ch["body"].split("\n\n")
        for p_text in paragraphs:
            p_text = p_text.strip()
            if not p_text:
                continue

            # Szenenumbruch
            if p_text in ("***", "---", "* * *"):
                scene_break = doc.add_paragraph("* * *")
                scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
                scene_break.space_before = Pt(18)
                scene_break.space_after = Pt(18)
                continue

            para = doc.add_paragraph(p_text)
            para.paragraph_format.first_line_indent = Cm(1)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = "Georgia"

        doc.add_page_break()

    doc.save(str(output_path))
    print(f"✅ DOCX erstellt: {output_path}")
    print(f"   {len(chapters)} Kapitel, {metadata['title']}")

if __name__ == "__main__":
    root = find_project_root()
    metadata = get_metadata(root)
    filename = metadata["title"].lower().replace(" ", "-") + ".docx"
    output = root / "_bmad-output" / "export" / filename
    output.parent.mkdir(parents=True, exist_ok=True)
    export_docx(root, output)
